"""브랜드 메시지 파일 텍스트 추출 — txt/docx/pdf/html 4종 지원.

Phase 0.6 BC-3 (pypdf → pdfplumber fallback) + BC-4 (python-docx) lessons 반영.
HTML 은 BeautifulSoup 으로 텍스트 추출 (스크립트/스타일/주석 제거).

본 모듈은 LLM 호출 0건 — 파일 → text 변환만 담당. 요약은 plan_generator 단계.
"""

from __future__ import annotations

import logging
import tempfile
from pathlib import Path

from domain.brand_card.model import BrandCardError

logger = logging.getLogger(__name__)


_SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".html", ".htm"}


class UnsupportedSourceError(BrandCardError):
    """지원하지 않는 파일 확장자 또는 인코딩."""


class SourceParseError(BrandCardError):
    """파싱 실패 — 손상된 파일, 권한 문제 등."""


def parse_source_file(path: Path) -> str:
    """파일 1개 → plain text. 확장자별 분기.

    Returns: 추출된 텍스트 (빈 문자열 가능 — 빈 파일이거나 텍스트 없는 PDF).

    Raises:
        UnsupportedSourceError: 지원 안 되는 확장자.
        SourceParseError: 읽기/파싱 실패.
        FileNotFoundError: 파일 미존재.
    """
    if not path.exists():
        raise FileNotFoundError(f"source file not found: {path}")
    suffix = path.suffix.lower()
    if suffix not in _SUPPORTED_EXTENSIONS:
        raise UnsupportedSourceError(
            f"확장자 {suffix!r} 미지원. 허용: {sorted(_SUPPORTED_EXTENSIONS)}"
        )
    if suffix in (".txt", ".md"):
        return _parse_text(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    return _parse_html(path)


def parse_source_bytes(suffix: str, data: bytes) -> str:
    """바이트 입력 → plain text. presigned 다운로드 결과 직접 처리용.

    docx/pdf 라이브러리는 파일 경로를 요구하는 경우가 많아 임시파일에 쓰고
    `parse_source_file` 에 위임한다. txt/md/html 만 BytesIO 직처리 가능하지만
    구현 단일화를 위해 모두 임시파일 경로로 처리한다.
    """
    suffix_lower = suffix.lower()
    if suffix_lower not in _SUPPORTED_EXTENSIONS:
        raise UnsupportedSourceError(
            f"확장자 {suffix_lower!r} 미지원. 허용: {sorted(_SUPPORTED_EXTENSIONS)}"
        )
    if not data:
        return ""

    with tempfile.NamedTemporaryFile(suffix=suffix_lower, delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        return parse_source_file(tmp_path)
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except OSError:
            logger.warning("source_parser.tmp_unlink_failed path=%s", tmp_path)


def _parse_text(path: Path) -> str:
    """UTF-8 우선, CP949 폴백 (한국어 텍스트 호환)."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="cp949")
        except UnicodeDecodeError as exc:
            raise SourceParseError(f"인코딩 추정 실패: {path}") from exc


def _parse_docx(path: Path) -> str:
    """python-docx 로 paragraph + 표 셀 평탄화 추출 (BC-4 lessons)."""
    try:
        from docx import Document
    except ImportError as exc:
        raise SourceParseError("python-docx 미설치 — pyproject 의존성 확인") from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise SourceParseError(f"docx 열기 실패: {path}") from exc

    chunks: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    chunks.append(cell_text)
    return "\n".join(chunks)


def _parse_pdf(path: Path) -> str:
    """BC-3 lesson — pypdf 우선, 빈 결과 시 pdfplumber fallback.

    스캔 PDF (텍스트 레이어 없음) 는 두 라이브러리 모두 빈 문자열 반환 가능.
    """
    text = _parse_pdf_with_pypdf(path)
    if text.strip():
        return text
    logger.info("pdf.pypdf_empty path=%s — fallback to pdfplumber", path)
    return _parse_pdf_with_pdfplumber(path)


def _parse_pdf_with_pypdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise SourceParseError("pypdf 미설치") from exc
    try:
        reader = PdfReader(str(path))
        chunks: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(page_text)
        return "\n\n".join(chunks)
    except Exception as exc:
        logger.warning("pdf.pypdf_failed path=%s err=%s", path, exc)
        return ""


def _parse_pdf_with_pdfplumber(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise SourceParseError("pdfplumber 미설치") from exc
    try:
        chunks: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    chunks.append(page_text)
        return "\n\n".join(chunks)
    except Exception as exc:
        raise SourceParseError(f"pdfplumber 실패: {path}") from exc


def _parse_html(path: Path) -> str:
    """BeautifulSoup 으로 텍스트 추출 — script/style/주석 제거."""
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise SourceParseError("beautifulsoup4 미설치") from exc
    try:
        raw = _parse_text(path)  # UTF-8/CP949 폴백 활용
    except SourceParseError:
        # HTML 은 종종 다른 인코딩 — bs4 가 직접 처리하도록
        raw = path.read_bytes().decode("utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)
